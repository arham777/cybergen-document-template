import streamlit as st
import tempfile
import os
import sys
from datetime import datetime
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE

# Import functions from cybergen_template.py
from cybergen_template import (
    extract_text_from_pdf,
    extract_text_from_docx,
    is_heading,
    add_space_after_paragraph,
    set_document_margins,
    add_current_date,
    format_paragraph
)

# Set page configuration
st.set_page_config(
    page_title="CyberGen Document Formatter",
    page_icon="📄",
    layout="wide"
)

# Custom CSS for better appearance
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1E3A8A;
        margin-bottom: 0.5rem;
    }
    .sub-header {
        font-size: 1.25rem;
        color: #6B7280;
        margin-bottom: 2rem;
    }
    .stButton button {
        background-color: #1E3A8A;
        color: white;
        font-weight: bold;
        padding: 0.5rem 1rem;
        border-radius: 0.25rem;
    }
    .info-box {
        background-color: #F3F4F6;
        padding: 1rem;
        border-radius: 0.5rem;
        margin-bottom: 1rem;
    }
</style>
""", unsafe_allow_html=True)

# Header section
st.markdown("<h1 class='main-header'>CyberGen Document Formatter</h1>", unsafe_allow_html=True)
st.markdown("<p class='sub-header'>Professional document formatting with automatic heading detection and styling</p>", unsafe_allow_html=True)

# Ensure the template exists
TEMPLATE_PATH = "cybergen-template.docx"

# Check if template exists and create if needed
if not os.path.exists(TEMPLATE_PATH):
    st.warning("Default template not found. Creating a basic template...")
    try:
        # Create a new document
        doc = docx.Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
        
        # Add a header
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = "CyberGen Document"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a footer
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "Confidential"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save the template
        doc.save(TEMPLATE_PATH)
        st.success("Basic template created successfully!")
    except Exception as e:
        st.error(f"Failed to create template: {str(e)}")

# Function to process document and create formatted output
def process_document(input_type, input_content, template_path=TEMPLATE_PATH):
    try:
        # Create a temporary output file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_output:
            output_filename = temp_output.name
        
        if input_type == "text":
            # Process text input
            doc = docx.Document(template_path)
            
            # Set margins
            set_document_margins(doc)
            
            # Add current date
            add_current_date(doc)
            
            # Process input text
            paragraphs = input_content.strip().split('\n')
            
            for paragraph_text in paragraphs:
                if not paragraph_text.strip():
                    continue
                    
                # Check if this paragraph is a heading
                heading_status = is_heading(paragraph_text)
                
                # Add paragraph with appropriate formatting
                p = doc.add_paragraph()
                run = p.add_run(paragraph_text)
                
                # Apply formatting based on heading status
                run.font.size = Pt(14) if heading_status else Pt(12.5)
                run.bold = True if heading_status else False
                run.underline = WD_UNDERLINE.SINGLE if heading_status else None
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if heading_status else WD_ALIGN_PARAGRAPH.JUSTIFY
                
                # Add proper spacing
                add_space_after_paragraph(p, is_heading=heading_status)
            
            # Apply widow/orphan control
            for paragraph in doc.paragraphs:
                paragraph.paragraph_format.widow_control = True
            
            # Save the document
            doc.save(output_filename)
            return output_filename
            
        elif input_type == "file":
            # Save temp file
            temp_input = None
            file_ext = os.path.splitext(input_content.name.lower())[1]
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
                temp_file.write(input_content.getvalue())
                temp_input = temp_file.name
            
            # Create new doc from template
            doc = docx.Document(template_path)
            
            # Set margins
            set_document_margins(doc)
            
            # Add current date
            add_current_date(doc)
            
            # Process based on file type
            if file_ext == '.pdf':
                # Extract text from PDF
                text_content = extract_text_from_pdf(temp_input)
                if text_content:
                    paragraphs = text_content.split('\n\n')
                    
                    for paragraph in paragraphs:
                        if paragraph.strip():
                            # Check if this paragraph is a heading
                            heading_status = is_heading(paragraph)
                            
                            # Add paragraph with appropriate formatting
                            p = doc.add_paragraph()
                            run = p.add_run(paragraph)
                            
                            # Format paragraph
                            run.font.size = Pt(14) if heading_status else Pt(12.5)
                            run.bold = True if heading_status else False
                            run.underline = WD_UNDERLINE.SINGLE if heading_status else None
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if heading_status else WD_ALIGN_PARAGRAPH.JUSTIFY
                            
                            # Add proper spacing
                            add_space_after_paragraph(p, is_heading=heading_status)
            
            elif file_ext in ('.docx', '.doc'):
                # Extract content from Word doc
                source_doc = docx.Document(temp_input)
                
                # Copy each paragraph
                for para in source_doc.paragraphs:
                    if para.text.strip():  # Skip empty paragraphs
                        # Check if this paragraph is a heading
                        heading_status = is_heading(para.text)
                        
                        # Add paragraph to template
                        new_para = doc.add_paragraph()
                        
                        # Copy text with formatting
                        for run in para.runs:
                            new_run = new_para.add_run(run.text)
                            # Copy run formatting
                            new_run.bold = run.bold if not heading_status else True
                            new_run.italic = run.italic
                            new_run.underline = run.underline if not heading_status else WD_UNDERLINE.SINGLE
                            # Set font size based on heading status
                            new_run.font.size = Pt(14) if heading_status else Pt(12.5)
                        
                        # Set alignment based on heading status
                        new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER if heading_status else WD_ALIGN_PARAGRAPH.JUSTIFY
                        
                        # Add proper spacing
                        add_space_after_paragraph(new_para, is_heading=heading_status)
                        
                        # If there are no runs (plain paragraph)
                        if not para.runs and para.text.strip():
                            new_run = new_para.add_run(para.text)
                            new_run.font.size = Pt(14) if heading_status else Pt(12.5)
                            new_run.bold = True if heading_status else False
                            new_run.underline = WD_UNDERLINE.SINGLE if heading_status else None
            
            # Apply widow/orphan control
            for paragraph in doc.paragraphs:
                paragraph.paragraph_format.widow_control = True
            
            # Save document
            doc.save(output_filename)
            
            # Clean up temp input file
            if temp_input and os.path.exists(temp_input):
                os.unlink(temp_input)
                
            return output_filename
    
    except Exception as e:
        st.error(f"Error processing document: {str(e)}")
        import traceback
        st.error(traceback.format_exc())
        return None

# Main UI
st.markdown("<div class='info-box'>", unsafe_allow_html=True)
st.markdown("""
### Key Features:
- Smart heading detection (ALL CAPS, colons, bullets/numbers)
- Professional formatting for headings and paragraphs
- Consistent spacing and alignment
- Pagination controls for better readability
- Current date automatically added at the top right
""")
st.markdown("</div>", unsafe_allow_html=True)

# Create tabs for different input methods
tab1, tab2 = st.tabs(["Enter Text", "Upload Document"])

with tab1:
    st.write("Enter your document text below:")
    text_input = st.text_area(
        "Document content will be formatted with professional styling",
        height=300,
        help="Text will be analyzed for headings and formatted accordingly"
    )
    
    if st.button("Generate Formatted Document", key="text_button"):
        if text_input:
            with st.spinner("Formatting document..."):
                result = process_document("text", text_input)
                
                if result:
                    with open(result, "rb") as file:
                        # Provide download button
                        st.download_button(
                            label="Download Formatted Document",
                            data=file,
                            file_name="formatted_document.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    
                    # Clean up temp file after download offered
                    try:
                        os.unlink(result)
                    except:
                        pass
                        
                    st.success("Document formatted successfully!")
        else:
            st.warning("Please enter some text first")

with tab2:
    st.write("Upload a Word document or PDF:")
    uploaded_file = st.file_uploader(
        "Choose a file",
        type=["docx", "pdf"],
        help="Your document will be processed and formatted according to our styling rules"
    )
    
    if st.button("Generate Formatted Document", key="file_button"):
        if uploaded_file is not None:
            with st.spinner("Formatting document..."):
                result = process_document("file", uploaded_file)
                
                if result:
                    with open(result, "rb") as file:
                        # Provide download button
                        st.download_button(
                            label="Download Formatted Document",
                            data=file,
                            file_name="formatted_document.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    
                    # Clean up temp file after download offered
                    try:
                        os.unlink(result)
                    except:
                        pass
                        
                    st.success("Document formatted successfully!")
        else:
            st.warning("Please upload a file first")

# App footer
st.markdown("---")
st.markdown(
    "Made with ❤️ | "
    f"Current Date: {datetime.now().strftime('%b %d, %Y')}"
)

# Sample text for testing
with st.expander("Sample Text for Testing"):
    st.code("""DOCUMENT TITLE

This is an example document to demonstrate the formatting capabilities
of the CyberGen Document Formatter.

First Section:

This text follows a heading with a colon and demonstrates how headings
are kept with their content across page breaks.

• Bulleted item as heading

This text demonstrates that bulleted items are detected as headings
and formatted accordingly.

Regular paragraph with more content. This shows the standard paragraph
formatting with justified text and appropriate spacing after paragraphs.
""", language="text") 