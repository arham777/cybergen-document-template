import streamlit as st
import os
import tempfile
from datetime import datetime
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE

# First, check if imports will work (this helps with deployment)
try:
    # Import just the essentials for our functionality
    import PyPDF2
    pdf_support = True
except ImportError:
    st.warning("PyPDF2 not available. PDF support will be disabled.")
    pdf_support = False

# Import main module, but handle potential errors
try:
    from cybergen_template import (
        extract_text_from_pdf,
        extract_text_from_docx,
        is_heading,
        add_space_after_paragraph,
        set_document_margins,
        add_current_date,
        format_paragraph,
        insert_text_into_template,
        copy_document_to_template
    )
    import_success = True
except ImportError as e:
    st.error(f"Error importing cybergen_template: {str(e)}")
    import_success = False

# Set page configuration
st.set_page_config(
    page_title="CyberGen Document Formatter",
    page_icon="📄",
    layout="wide"
)

# Custom CSS for better appearance
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        color: #1E3A8A;
        margin-bottom: 0.5rem;
    }
    .sub-header {
        font-size: 1.25rem;
        color: #6B7280;
        margin-bottom: 2rem;
    }
    .stButton button {
        background-color: #1E3A8A;
        color: white;
        font-weight: bold;
        padding: 0.5rem 1rem;
        border-radius: 0.25rem;
    }
    .info-box {
        background-color: #F3F4F6;
        padding: 1rem;
        border-radius: 0.5rem;
        margin-bottom: 1rem;
    }
    .template-box {
        border: 1px solid #E5E7EB;
        padding: 1rem;
        border-radius: 0.25rem;
        background-color: #F9FAFB;
        margin-bottom: 1rem;
    }
</style>
""", unsafe_allow_html=True)

# Header section
st.markdown("<h1 class='main-header'>CyberGen Document Formatter</h1>", unsafe_allow_html=True)
st.markdown("<p class='sub-header'>Professional document formatting with automatic heading detection and styling</p>", unsafe_allow_html=True)

# Initialize session state for template storage
if 'template_path' not in st.session_state:
    st.session_state.template_path = "cybergen-template.docx"  # Default template
if 'custom_template_uploaded' not in st.session_state:
    st.session_state.custom_template_uploaded = False
if 'template_info' not in st.session_state:
    st.session_state.template_info = "Using default template"

# Function to create a basic template
def create_basic_template(template_path):
    try:
        # Create a new document
        doc = docx.Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
        
        # Add a header
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = "CyberGen Document"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add a footer
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "Confidential"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save the template
        doc.save(template_path)
        return True
    except Exception as e:
        st.error(f"Failed to create template: {str(e)}")
        return False

# Ensure the default template exists
if not os.path.exists(st.session_state.template_path) and not st.session_state.custom_template_uploaded:
    st.warning("Default template not found. Creating a basic template...")
    if create_basic_template(st.session_state.template_path):
        st.success("Basic template created successfully!")

# Template configuration section in sidebar
with st.sidebar:
    st.header("Document Template")
    
    st.markdown("<div class='template-box'>", unsafe_allow_html=True)
    
    # Option to upload a custom template
    uploaded_template = st.file_uploader(
        "Upload your own template document",
        type=["docx"],
        help="This will be used as the base template for your document formatting"
    )
    
    if uploaded_template:
        # Create a temporary file for the uploaded template
        temp_template_path = f"uploaded_template_{uploaded_template.name}"
        with open(temp_template_path, "wb") as f:
            f.write(uploaded_template.getvalue())
        
        # Update session state
        st.session_state.template_path = temp_template_path
        st.session_state.custom_template_uploaded = True
        st.session_state.template_info = f"Using custom template: {uploaded_template.name}"
        
        # Show template details
        st.success(f"Custom template '{uploaded_template.name}' loaded successfully!")
        
        # Option to revert to default template
        if st.button("Use Default Template Instead"):
            if os.path.exists(temp_template_path):
                os.remove(temp_template_path)
            st.session_state.template_path = "cybergen-template.docx"
            st.session_state.custom_template_uploaded = False
            st.session_state.template_info = "Using default template"
            st.experimental_rerun()
    else:
        st.info("Using the default template. Upload your own template for customized formatting.")
        
        # Create default template if needed
        if not os.path.exists("cybergen-template.docx"):
            if st.button("Create Default Template"):
                if create_basic_template("cybergen-template.docx"):
                    st.success("Default template created successfully!")
                    st.experimental_rerun()
    
    st.markdown("</div>", unsafe_allow_html=True)
    
    # Additional options
    st.markdown("### Options")
    st.markdown("Current template: " + st.session_state.template_info)

# Display error if imports failed
if not import_success:
    st.error("Cannot continue due to import errors. Please check the logs.")
    st.stop()

# Function to process document and create formatted output
def process_document(input_type, input_content, template_path=st.session_state.template_path):
    try:
        # Create a temporary output file
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_output:
            output_filename = temp_output.name
        
        if input_type == "text":
            # Process text using the existing function
            result = insert_text_into_template(
                input_content, 
                template_path=template_path,
                output_filename=output_filename
            )
            return result
            
        elif input_type == "file":
            # Save temp file
            temp_input = None
            file_ext = os.path.splitext(input_content.name.lower())[1]
            
            with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as temp_file:
                temp_file.write(input_content.getvalue())
                temp_input = temp_file.name
            
            # Check if PDF is supported
            if file_ext == '.pdf' and not pdf_support:
                st.error("PDF support is not available in this deployment.")
                if temp_input and os.path.exists(temp_input):
                    os.unlink(temp_input)
                return None
            
            # Use the existing function to copy document to template
            result = copy_document_to_template(
                temp_input,
                template_path=template_path,
                output_filename=output_filename
            )
            
            # Clean up temp input file
            if temp_input and os.path.exists(temp_input):
                os.unlink(temp_input)
                
            return result
    
    except Exception as e:
        st.error(f"Error processing document: {str(e)}")
        import traceback
        st.error(traceback.format_exc())
        return None

# Main UI
st.markdown("<div class='info-box'>", unsafe_allow_html=True)
st.markdown("""
### Key Features:
- Smart heading detection (ALL CAPS, colons, bullets/numbers)
- Professional formatting for headings and paragraphs
- Consistent spacing and alignment
- Pagination controls for better readability
- Current date automatically added at the top right
""")
st.markdown("</div>", unsafe_allow_html=True)

# Create tabs for different input methods
tab1, tab2 = st.tabs(["Enter Text", "Upload Document"])

with tab1:
    st.write("Enter your document text below:")
    text_input = st.text_area(
        "Document content will be formatted with professional styling",
        height=300,
        help="Text will be analyzed for headings and formatted accordingly"
    )
    
    output_name = st.text_input(
        "Output filename:", 
        value="formatted_document.docx",
        help="Name of the output document file (will be appended with .docx if not included)"
    )
    
    # Ensure output filename has .docx extension
    if not output_name.lower().endswith('.docx'):
        output_name += '.docx'
    
    if st.button("Generate Formatted Document", key="text_button"):
        if text_input:
            with st.spinner(f"Formatting document using {st.session_state.template_info}..."):
                result = process_document("text", text_input)
                
                if result:
                    with open(result, "rb") as file:
                        # Provide download button
                        st.download_button(
                            label="Download Formatted Document",
                            data=file,
                            file_name=output_name,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    
                    # Clean up temp file after download offered
                    try:
                        os.unlink(result)
                    except:
                        pass
                        
                    st.success(f"Document formatted successfully using {st.session_state.template_info}!")
        else:
            st.warning("Please enter some text first")

with tab2:
    st.markdown(f"**Template: {st.session_state.template_info}**")
    
    file_types = ["docx", "pdf"] if pdf_support else ["docx"]
    file_type_help = "Upload a Word document or PDF" if pdf_support else "Upload a Word document (PDF support not available in this deployment)"
    
    st.write(file_type_help)
    uploaded_file = st.file_uploader(
        "Choose a file to format",
        type=file_types,
        help="Your document will be processed and formatted according to our styling rules using the selected template"
    )
    
    if uploaded_file is not None:
        # Display file info
        col1, col2 = st.columns(2)
        with col1:
            st.write("File name:", uploaded_file.name)
        with col2:
            st.write("File size:", f"{uploaded_file.size / 1024:.1f} KB")
    
    output_name = st.text_input(
        "Output filename:", 
        value="formatted_document.docx", 
        key="file_output_name",
        help="Name of the output document file (will be appended with .docx if not included)"
    )
    
    # Ensure output filename has .docx extension
    if not output_name.lower().endswith('.docx'):
        output_name += '.docx'
    
    if st.button("Generate Formatted Document", key="file_button"):
        if uploaded_file is not None:
            with st.spinner(f"Formatting document using {st.session_state.template_info}..."):
                result = process_document("file", uploaded_file)
                
                if result:
                    with open(result, "rb") as file:
                        # Provide download button
                        st.download_button(
                            label="Download Formatted Document",
                            data=file,
                            file_name=output_name,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    
                    # Clean up temp file after download offered
                    try:
                        os.unlink(result)
                    except:
                        pass
                        
                    st.success(f"Document formatted successfully using {st.session_state.template_info}!")
        else:
            st.warning("Please upload a file first")

# App footer
st.markdown("---")
st.markdown(
    "Made with ❤️ | "
    f"Current Date: {datetime.now().strftime('%b %d, %Y')}"
)

# Sample text for testing
with st.expander("Sample Text for Testing"):
    st.code("""DOCUMENT TITLE

This is an example document to demonstrate the formatting capabilities
of the CyberGen Document Formatter.

First Section:

This text follows a heading with a colon and demonstrates how headings
are kept with their content across page breaks.

• Bulleted item as heading

This text demonstrates that bulleted items are detected as headings
and formatted accordingly.

Regular paragraph with more content. This shows the standard paragraph
formatting with justified text and appropriate spacing after paragraphs.
""", language="text") 