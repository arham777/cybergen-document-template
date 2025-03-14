import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_template(template_name="cybergen-template.docx"):
    """
    Creates a simple template document for the CyberGen Document Formatter POC.
    
    Args:
        template_name: Name for the template file
    
    Returns:
        Path to the created template
    """
    try:
        # Create a new document
        doc = docx.Document()
        
        # Set margins
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
        
        # Add a header (optional)
        header = doc.sections[0].header
        header_para = header.paragraphs[0]
        header_para.text = "CyberGen Document"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in header_para.runs:
            run.font.size = Pt(10)
            
        # Add a footer (optional)
        footer = doc.sections[0].footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "Confidential"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer_para.runs:
            run.font.size = Pt(10)
        
        # Save the template
        doc.save(template_name)
        return os.path.abspath(template_name)
        
    except Exception as e:
        print(f"Error creating template: {str(e)}")
        return None

if __name__ == "__main__":
    template_path = create_template()
    if template_path:
        print(f"Template created successfully at: {template_path}")
        print("You can now use this template with the cybergen_poc.py script.")
    else:
        print("Failed to create template. Check error messages above.") 