import docx
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE, WD_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION
import os
import sys
import re
from datetime import datetime
import PyPDF2  # For PDF text extraction

def extract_text_from_pdf(file_path):
    """
    Extract text content from a PDF file.
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        str: Extracted text content
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if not file_path.lower().endswith('.pdf'):
            raise ValueError("File must be a PDF")
        
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text() + "\n\n"
        
        return text.strip()
    
    except Exception as e:
        print(f"Error extracting text from PDF: {str(e)}")
        return None

def create_document_from_template(input_text, output_filename="generated_document.docx"):
    """
    Creates a document based on the CyberGen template using the provided input text.
    
    Args:
        input_text (str): The text content to be formatted according to the template
        output_filename (str): The name of the output document file
    
    Returns:
        str: Path to the created document
    """
    try:
        # Create a new document
        doc = docx.Document()
        
        # Set page margins (1 inch on all sides)
        section = doc.sections[0]
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
        # Set default font for the document
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12.5)
        
        # Add header with logo
        header = section.header
        
        # Create a table for the header content
        header_table = header.add_table(1, 2, width=Inches(6.5))
        
        # Left column: Add wavy line design
        left_cell = header_table.cell(0, 0)
        left_para = left_cell.paragraphs[0]
        left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        left_run = left_para.add_run("≈≈≈≈≈≈≈≈≈")
        left_run.font.size = Pt(28)
        left_run.font.color.rgb = RGBColor(173, 216, 230)  # Light blue color
        
        # Right column: Add CYBERGEN logo text
        right_cell = header_table.cell(0, 1)
        logo_para = right_cell.paragraphs[0]
        logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add CYBERGEN logo text with "One Team" subtitle
        logo_run = logo_para.add_run("CYBERGEN")
        logo_run.bold = True
        logo_run.font.size = Pt(28)
        logo_run.font.color.rgb = RGBColor(0, 174, 239)  # CYBERGEN blue
        
        # Add "One Team" in smaller text
        logo_para.add_run("\n")
        team_run = logo_para.add_run("One Team")
        team_run.font.size = Pt(14)
        team_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color
        
        # Add date to the document (right-aligned, below logo)
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        current_date = datetime.now().strftime("%b %d, %Y")  # Format: Nov 28, 2024
        date_run = date_para.add_run(current_date)
        date_run.font.size = Pt(13)
        date_run.bold = True  # Make date bold
        date_para.space_after = Pt(12)
        
        # Process the input text
        paragraphs = input_text.strip().split('\n')
        
        # Check if the document has content
        if not paragraphs:
            raise ValueError("Input text is empty. Please provide some content.")
        
        # Add standard "TO WHOM IT MAY CONCERN" header
        concern_para = doc.add_paragraph()
        concern_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        concern_run = concern_para.add_run("TO WHOM IT MAY CONCERN")
        concern_run.bold = True
        concern_run.underline = WD_UNDERLINE.SINGLE  # Add underline to header
        concern_run.font.size = Pt(14)
        concern_para.space_after = Pt(12)
        
        # Extract and format header if it exists (enclosed in ** and with .underline)
        header_pattern = r'\*\*\[(.*?)\]\{\.underline\}\*\*'
        
        for i, paragraph in enumerate(paragraphs):
            # Check if this paragraph is a header
            header_match = re.search(header_pattern, paragraph)
            if header_match:
                # Add the header with proper formatting
                header_text = header_match.group(1)
                header = doc.add_paragraph()
                header_run = header.add_run(header_text)
                header_run.bold = True
                header_run.underline = WD_UNDERLINE.SINGLE
                header_run.font.size = Pt(14)
                header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                header.space_after = Pt(12)  # Add spacing after header
                
                # Remove this paragraph from the list so it's not processed again
                paragraphs[i] = ""
        
        # Process remaining paragraphs as body text
        for paragraph in paragraphs:
            if paragraph.strip():  # Skip empty paragraphs
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Change to justified alignment
                p.space_after = Pt(6)  # Add spacing between paragraphs
                run = p.add_run(paragraph)
                run.font.size = Pt(12.5)
        
        # Add signature section
        doc.add_paragraph().space_after = Pt(24)  # Add space before signature
        
        sincerely = doc.add_paragraph()
        sincerely_run = sincerely.add_run("Sincerely,")
        sincerely_run.bold = True
        sincerely_run.font.size = Pt(12.5)
        sincerely.space_after = Pt(36)  # Space for signature
        
        name = doc.add_paragraph()
        name.add_run("Gazala Arif").font.size = Pt(12.5)
        name.add_run().add_break()
        name.add_run("HR Manager").font.size = Pt(12.5)
        
        # Add footer with contact information
        footer = section.footer
        footer_table = footer.add_table(1, 3, width=Inches(8))
        
        # Left column: Phone
        phone_cell = footer_table.cell(0, 0)
        phone_para = phone_cell.paragraphs[0]
        phone_icon = phone_para.add_run("📞 ")
        phone_icon.font.size = Pt(12)
        phone_icon.font.color.rgb = RGBColor(0, 174, 239)  # Light blue
        phone_para.add_run("Phone\n").bold = True
        phone_text = phone_para.add_run("+92-41-5487208")
        phone_text.font.size = Pt(12.5)
        
        # Middle column: Website
        web_cell = footer_table.cell(0, 1)
        web_para = web_cell.paragraphs[0]
        web_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        web_icon = web_para.add_run("🌐 ")
        web_icon.font.size = Pt(12)
        web_icon.font.color.rgb = RGBColor(0, 174, 239)  # Light blue
        web_para.add_run("Website\n").bold = True
        web_text = web_para.add_run("www.cybergen.com")
        web_text.font.size = Pt(12.5)
        
        # Right column: Location
        loc_cell = footer_table.cell(0, 2)
        loc_para = loc_cell.paragraphs[0]
        loc_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Location text
        loc_icon = loc_para.add_run("📍 ")
        loc_icon.font.size = Pt(12)
        loc_icon.font.color.rgb = RGBColor(0, 174, 239)  # Light blue
        loc_para.add_run("Location\n").bold = True
        loc_text = loc_para.add_run("Office 107, Mall of Faisalabad, Faisalabad")
        loc_text.font.size = Pt(12.5)
        
        # Save the document
        try:
            doc.save(output_filename)
            return os.path.abspath(output_filename)
        except Exception as e:
            raise IOError(f"Failed to save document: {str(e)}")
    
    except Exception as e:
        print(f"Error creating document: {str(e)}")
        return None


def parse_document(file_path):
    """
    Parse an existing document and return its text content.
    
    Args:
        file_path (str): Path to the document file
    
    Returns:
        str: The text content of the document
    """
    try:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check file extension to determine parsing method
        file_ext = os.path.splitext(file_path.lower())[1]
        
        if file_ext == '.pdf':
            # Parse PDF file
            return extract_text_from_pdf(file_path)
        elif file_ext in ('.docx', '.doc'):
            # Parse Word document
        doc = docx.Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip():  # Only add non-empty paragraphs
                full_text.append(para.text)
        
        return '\n'.join(full_text)
        else:
            raise ValueError("File must be a Word document (.doc or .docx) or a PDF (.pdf)")
    
    except Exception as e:
        print(f"Error parsing document: {str(e)}")
        return None


def main():
    """
    Main function to handle user interaction and document processing.
    """
    print("CyberGen Document Formatter")
    print("==========================")
    
    while True:
        print("\nOptions:")
        print("1. Enter text directly")
        print("2. Import text from an existing document (Word or PDF)")
        print("3. Exit")
        
        choice = input("\nEnter your choice (1-3): ")
        
        if choice == '1':
            print("\nEnter your document text (type 'END' on a new line when finished):")
            lines = []
            while True:
                line = input()
                if line == 'END':
                    break
                lines.append(line)
            
            input_text = '\n'.join(lines)
            output_name = input("\nEnter output filename (leave blank for default 'generated_document.docx'): ")
            if not output_name:
                output_name = "generated_document.docx"
            elif not output_name.lower().endswith('.docx'):
                output_name += '.docx'
                
            document_path = create_document_from_template(input_text, output_name)
            if document_path:
                print(f"\nDocument successfully created at: {document_path}")
            
        elif choice == '2':
            file_path = input("\nEnter the path to the document (Word or PDF): ")
            document_text = parse_document(file_path)
            
            if document_text:
                print("\nDocument content:")
                print("----------------")
                print(document_text)
                print("----------------")
                
                proceed = input("\nDo you want to format this text using the CyberGen template? (y/n): ")
                if proceed.lower() == 'y':
                    output_name = input("\nEnter output filename (leave blank for default 'generated_document.docx'): ")
                    if not output_name:
                        output_name = "generated_document.docx"
                    elif not output_name.lower().endswith('.docx'):
                        output_name += '.docx'
                    
                    document_path = create_document_from_template(document_text, output_name)
                    if document_path:
                        print(f"\nDocument successfully created at: {document_path}")
        
        elif choice == '3':
            print("\nExiting program. Goodbye!")
            break
        
        else:
            print("\nInvalid choice. Please try again.")


if __name__ == "__main__":
    main()