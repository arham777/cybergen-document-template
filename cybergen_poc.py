import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
import os
from datetime import datetime

def poc_document_formatter(input_text, template_path="cybergen-template.docx", output_filename="poc_document.docx"):
    """
    POC document formatter that demonstrates key functionality:
    - Uses a template document
    - Adds current date at top right
    - Detects and formats headings
    - Applies consistent paragraph formatting
    - Ensures headings stay with content
    """
    try:
        # Check if template exists
        if not os.path.exists(template_path):
            print(f"Template not found: {template_path}")
            return None
            
        # Create document from template
        doc = docx.Document(template_path)
        
        # Set page margins
        for section in doc.sections:
            section.top_margin = Inches(1.5)
            section.bottom_margin = Inches(1.5)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)
        
        # Add current date
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = date_para.add_run(datetime.now().strftime("%b %d, %Y"))
        date_run.font.size = Pt(12)
        date_run.bold = True
        date_para.paragraph_format.space_after = Pt(12)
        
        # Process input text
        paragraphs = input_text.strip().split('\n')
        
        # Process each paragraph
        for paragraph_text in paragraphs:
            if not paragraph_text.strip():
                continue
                
            # Detect if this is a heading
            is_heading = False
            if len(paragraph_text.strip()) < 100:
                if paragraph_text.isupper() or paragraph_text.endswith(':'):
                    is_heading = True
                if (paragraph_text.strip().startswith(('•', '-', '*')) or 
                    (paragraph_text[0].isdigit() and '.' in paragraph_text[:3])):
                    is_heading = True
            
            # Create paragraph with appropriate formatting
            p = doc.add_paragraph()
            run = p.add_run(paragraph_text)
            
            if is_heading:
                # Heading formatting
                run.font.size = Pt(14)
                run.bold = True
                run.underline = WD_UNDERLINE.SINGLE
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_after = Pt(18)
                p.paragraph_format.keep_with_next = True
            else:
                # Normal paragraph formatting
                run.font.size = Pt(12.5)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(12)
        
        # Apply widow/orphan control
        for paragraph in doc.paragraphs:
            paragraph.paragraph_format.widow_control = True
            
        # Save the document
        doc.save(output_filename)
        return os.path.abspath(output_filename)
        
    except Exception as e:
        print(f"Error in POC document formatter: {str(e)}")
        return None

def main():
    """
    Simple POC interface to demonstrate document formatting
    """
    print("CyberGen Document Formatter - POC")
    print("=================================")
    
    # Get template path
    template = input("Enter template path (or press Enter for default 'cybergen-template.docx'): ")
    template = template if template else "cybergen-template.docx"
    
    # Get input text
    print("\nEnter your document text (type 'END' on a new line when finished):")
    print("Note: Text will be formatted according to these rules:")
    print("- Headings (ALL CAPS, ending with colon, or bullets/numbers) will be centered, bold, underlined")
    print("- Normal paragraphs will be justified with 12.5pt font size")
    lines = []
    while True:
        line = input()
        if line == 'END':
            break
        lines.append(line)
    
    input_text = '\n'.join(lines)
    
    # Get output filename
    output = input("\nEnter output filename (or press Enter for 'poc_document.docx'): ")
    output = output if output else "poc_document.docx"
    
    # Process document
    result = poc_document_formatter(input_text, template, output)
    
    if result:
        print(f"\nPOC document successfully created at: {result}")
        print("Key features demonstrated:")
        print("- Template-based document creation")
        print("- Automatic date insertion")
        print("- Heading detection and formatting")
        print("- Paragraph formatting with proper spacing")
        print("- Pagination control to keep headings with content")
    else:
        print("\nFailed to create POC document. Check error messages above.")

if __name__ == "__main__":
    main() 